import docx
#打开文件
doc=docx.Document('/Users/cucuser/Desktop/人间失格.docx')
#doc=docx.Document('\\Users\\cucuser\\Desktop\\人间失格.docx')
print(dir(doc)) #打印doc的所有可操作函数方法

#paragraphs函数
print(doc.paragraphs)   #打印文档中段落对象
print(len(doc.paragraphs))  #该段文本中段落数量
print(doc.paragraphs[0].text)   #该段中文本的字符串内容（没有样式信息）
for i in range(0,4):
    print(doc.paragraphs[i].text)#遍历各段内容

#run函数
print(doc.paragraphs[0].runs)   #获取对应段落中的runs个数
print(doc.paragraphs[0].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)

#添加一个tab
doc.paragraphs[0].runs[3].add_tab()
#添加文字
doc.paragraphs[0].runs[3].add_text('新添加一部分内容')
#-----设置字符样式------
#加粗文字
doc.paragraphs[0].runs[3].bold=True
#英文字母大写
doc.paragraphs[0].runs[3].font.all_caps=True
#换字体颜色
from docx.shared import RGBColor
doc.paragraphs[0].runs[3].font.color.rgb= RGBColor(0,233,233)
#定义字体
from docx.shared import Pt
doc.paragraphs[0].runs[0].font.size=Pt(18)
from docx.shared import Cm
doc.paragraphs[0].runs[6].font.size=Cm(1)

#定义字体
from docx.oxml.ns import qn
doc.paragraphs[0].runs[0].font.name='黑体'
doc.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsi'),'宋体')

#清除内容
doc.paragraphs[0].runs[2].clear()
#保存文件
doc.save('/Users/cucuser/Desktop/人间失格.docx')

