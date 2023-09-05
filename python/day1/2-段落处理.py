import docx

# 打开文件
doc = docx.Document('/Users/cucuser/Desktop/人间失格.docx')
#修改段落对齐
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
doc.paragraphs[1].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.paragraphs[2].alignment=WD_ALIGN_PARAGRAPH.RIGHT
doc.paragraphs[3].alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
doc.paragraphs[4].alignment=WD_ALIGN_PARAGRAPH.DISTRIBUTE
#插入新的一段，样式相同
doc.paragraphs[0].insert_paragraph_before
#保存文件
doc.save('/Users/cucuser/Desktop/人间失格.docx')

